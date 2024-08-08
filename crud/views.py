from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.apps import apps
from django.shortcuts import get_object_or_404
from django.urls import reverse_lazy
from . import forms
from .forms import ClientFormSet, InvoiceFormSet
from .models import *
from .serializers import *
from rest_framework import serializers
from rest_framework import viewsets
from rest_framework.views import APIView
from rest_framework import status
from rest_framework.response import Response
from rest_framework.pagination import PageNumberPagination
from django.db.models import Q
from django.core.exceptions import ImproperlyConfigured
from django.shortcuts import render,redirect
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
# Create your views here.
from django.http import HttpResponse
import importlib
from rest_framework.exceptions import NotFound
from django.conf import settings
from django.http import FileResponse
import os
from django.http import Http404
# imports for docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from .filters import *
def format_date(datetime_str):
    try:
        dt = datetime.strptime(datetime_str, '%Y-%m-%d %H:%M:%S')
        return dt.strftime('%Y-%m-%d')
    except ValueError:
        return datetime_str

# Function to set the background color of a cell
def set_cell_background(cell, color):
    # Get the tc element of the cell
    tc = cell._element
    # Create a new shading element
    tc_pr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)  # Set the fill attribute to the desired color
    tc_pr.append(shading)
def set_font_style(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = font_size
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
def add_styling(run,font_size=22):
    # Set the desired style
    run.font.size = Pt(22)
    run.font.name = 'Roboto'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Roboto')
# Mapping model names to their respective filter classes
FILTER_MAPPING = {
    'client': ClientFilter,
    'solarpanel': SolarPanelFilter,
    'inverter': InverterFilter,
    'structure': StructureFilter,
    'cabling': CablingFilter,
    'netmetering': NetMeteringFilter,
    'batteries': BatteriesFilter,
    'lightningarrestor': LightningArrestorFilter,
    'installation': InstallationFilter,
    'invoice': InvoiceFilter,
    'expenditures': ExpendituresFilter,
}
class DynamicModelViewSet(viewsets.ModelViewSet):

    def apply_filters(self, queryset):
        search_query = self.request.GET.get('q')
        date_range = self.request.GET.get('date_range')
        if search_query:
            query = Q()
            for field in self.get_searchable_fields():
                query |= Q(**{f"{field}__icontains": search_query})
            queryset = queryset.filter(query)
        if date_range:
            try:
                start_date, end_date = date_range.split(' to ')
                start_date = datetime.strptime(start_date, '%Y-%m-%d')
                end_date = datetime.strptime(end_date, '%Y-%m-%d')
                queryset = queryset.filter(created_at__range=(start_date, end_date))
            except ValueError:
                pass
        return queryset

    def get_searchable_fields(self):
        fields = []
        for field in self.model._meta.fields:
            if isinstance(field, (models.CharField, models.TextField)):
                fields.append(field.name)
            elif isinstance(field, models.ForeignKey):
                if hasattr(field.related_model, 'name'):
                    fields.append(f"{field.name}__name")
                elif hasattr(field.related_model, 'brand'):
                    fields.append(f"{field.name}__brand")
        return fields

    def apply_sort(self, queryset: models.QuerySet):
        sort_by = self.request.GET.get('sort_by', None)
        if sort_by:
            if sort_by.startswith('-'):
                sort_by = sort_by[1:]
            else:
                sort_by = f'-{sort_by}'
            queryset = queryset.order_by(sort_by)
        return queryset
    def get_queryset(self,request=None):
        model_name = self.kwargs['model_name']
        self.model = apps.get_model('crud', model_name)
        queryset = self.model.objects.all()
        filter_class = FILTER_MAPPING.get(model_name.lower())
        if filter_class and request:
            my_filter = filter_class(request.GET,queryset=queryset)
            queryset = my_filter.qs
        else:
            my_filter = None

        queryset = self.apply_filters(queryset)
        queryset = self.apply_sort(queryset)

        return queryset

    def get_serializer_class(self):
        model_name = self.kwargs['model_name']
        module_path = 'crud.serializers'
        serializer_class_name = model_name + 'Serializer'

        try:
            serializers_module = importlib.import_module(module_path)
            serializer_class = getattr(serializers_module, serializer_class_name)
        except (ImportError, AttributeError):
            raise NotFound(f"Serializer for model '{model_name}' not found.")
        
        return serializer_class
    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset(request)

        page = self.paginate_queryset(queryset)

        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
import tempfile
def index(request):
    file_path = os.path.join(settings.BASE_DIR, 'test.rest')  # Update this to the path of your file

    if os.path.exists(file_path):
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename='test.rest')  # 'as_attachment=True' makes it a download
    else:
        raise Http404("File not found.")

def manage_items(request):
    if request.method == 'POST':
        formset = ClientFormSet(request.POST, queryset=Client.objects.all())
        if formset.is_valid():
            formset.save()
            return redirect('success_url')  # Replace with your success URL
    else:
        formset = ClientFormSet(queryset=Client.objects.all())

    return render(request, 'manage_clients.html', {'formset': formset})
def manage_invoices(request):
    if request.method == 'POST':
        formset = InvoiceFormSet(request.POST)
        if formset.is_valid():
            formset.save()
            return redirect('success_url')  # Replace with your success URL
    else:
        formset = InvoiceFormSet(queryset=Invoice.objects.all())

    return render(request, 'manage_invoices.html', {'formset': formset})
def modify_and_send_file(request, invoice_id):
    try:
        invoice = get_object_or_404(Invoice, pk=invoice_id)
        partially_paid_records = PartiallyPaid.objects.filter(invoice=invoice_id)
        print('Partially Paid Records:', partially_paid_records)

        # Initialize an empty list to store the records
        installments_list = []
        partially_paid = 0
        # Iterate over the records and store them in the list as dictionaries
        for record in partially_paid_records:
            partially_paid += float(record.amount)
            record_dict = {
                "client_name": record.client.name,
                "amount": float(record.amount),
                "time": record.time.strftime('%Y-%m-%d %H:%M:%S')  # Format datetime as string
            }
            installments_list.append(record_dict)
                # Create a function to display foreign key relationships properly
        def get_foreign_key_display(instance, field_name):
            fk_instance = getattr(instance, field_name)
            return str(fk_instance) if fk_instance else 'None'

        # Print the inner content of the invoice object
        invoice_details = {
            'ID': invoice.id,
            "system_capacity": invoice.system_capacity,
            'Client Name': get_foreign_key_display(invoice, 'name'),
            'Solar Panel': get_foreign_key_display(invoice, 'solar_panel'),
            'Solar Panel Quantity': invoice.solar_panel_quantity,
            'Solar Panel Price': invoice.solar_panel_price,
            'Inverter': get_foreign_key_display(invoice, 'inverter'),
            'Inverter Quantity': invoice.inverter_quantity,
            'Inverter Price': invoice.inverter_price,
            'Structure': get_foreign_key_display(invoice, 'structure'),
            'Structure Quantity': invoice.structure_quantity,
            'Structure Price': invoice.structure_price,
            'Cabling': get_foreign_key_display(invoice, 'cabling'),
            'Cabling Quantity': invoice.cabling_quantity,
            'Cabling Price': invoice.cabling_price,
            'Net Metering': get_foreign_key_display(invoice, 'net_metering'),
            'Net Metering Quantity': invoice.net_metering_quantity,
            'Net Metering Price': invoice.net_metering_price,
            'Battery': get_foreign_key_display(invoice, 'battery'),
            'Battery Quantity': invoice.battery_quantity,
            'Battery Price': invoice.battery_price,
            'Lightning Arrestor': get_foreign_key_display(invoice, 'lightning_arrestor'),
            'Lightning Arrestor Quantity': invoice.lightning_arrestor_quantity,
            'Lightning Arrestor Price': invoice.lightning_arrestor_price,
            'Installation': get_foreign_key_display(invoice, 'installation'),
            'Installation Quantity': invoice.installation_quantity,
            'Installation Price': invoice.installation_price,
            'Discount': invoice.discount,
            'Shipping Charges': invoice.shipping_charges,
            'Status': invoice.status,
            'Amount Paid': invoice.amount_paid,
            'Total': invoice.total,
            'Created At': invoice.created_at,
            'Updated At': invoice.updated_at,
        }

        # Print invoice details
        for key, value in invoice_details.items():
            print(f"{key}: {value}")
        print('Invoice Details:', invoice.__dict__)
        client_data = [
            {'name': '' + invoice.solar_panel.name, 'quantity': invoice.solar_panel_quantity, 'price': invoice.solar_panel_price},
            {'name': '' + invoice.inverter.name, 'quantity': invoice.inverter_quantity, 'price': invoice.inverter_price},
            {'name': '' + invoice.structure.name, 'quantity': invoice.structure_quantity, 'price': invoice.structure_price},
            {'name': '' + invoice.cabling.name, 'quantity': invoice.cabling_quantity, 'price': invoice.cabling_price},
            {'name': '' + invoice.net_metering.name, 'quantity': invoice.net_metering_quantity, 'price': invoice.net_metering_price},
            {'name': '' + invoice.battery.name, 'quantity': invoice.battery_quantity, 'price': invoice.battery_price},
            {'name': '' + invoice.lightning_arrestor.name, 'quantity': invoice.lightning_arrestor_quantity, 'price': invoice.lightning_arrestor_price},
            {'name': '' + invoice.installation.name, 'quantity': invoice.installation_quantity, 'price': invoice.installation_price},
            ]
        
        final_block_data = [
        ("Subtotal", invoice.total),
        ("Discount", invoice.discount),
        ("Discounted Amount", invoice.total - int(invoice.discount)),
        ("Partially Paid", partially_paid),
        ("Shipping", invoice.shipping_charges),
        ("Total", invoice.total - int(invoice.discount) + int(invoice.shipping_charges) - int(invoice.amount_paid))
        ]
        header_path = os.path.join(settings.MEDIA_ROOT, 'invoices', 'header.docx')
        file_path = os.path.join(settings.MEDIA_ROOT, 'invoices', 'template.docx')
        footer_path = os.path.join(settings.MEDIA_ROOT, 'invoices', 'footer.docx')
        if not os.path.exists(file_path):
            raise Http404("File not found.")
    except Invoice.DoesNotExist:
        raise Http404("Invoice not found.")
    def list_get(lst, index, default=None):
        try:
            return lst[index]
        except IndexError:
            return default
    def handle_table0(table):
        # Set Status
        
        row = table.rows[0].cells
        cell = row[0]
        # Set text properties
        paragraph = cell.paragraphs[0]
        paragraph.clear()  # Clear existing text in the paragraph
        run = paragraph.add_run(str(invoice.status))
        run.bold = True
        run.font.name = 'Roboto'

        # This part ensures the font name is set correctly
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Roboto')
        
        run.font.size = Pt(24)  # Adjust the size as needed
        run.font.color.rgb = RGBColor(127, 127, 127)  # Dark blue color
        # Set date
        row = table.rows[1].cells
        cell = row[4]
        # Set text properties
        paragraph = cell.paragraphs[0]
        paragraph.clear()  # Clear existing text in the paragraph
        run = paragraph.add_run(str(invoice.updated_at.strftime('%Y-%m-%d')))
        run.bold = True
        run.font.name = 'Roboto'

        # This part ensures the font name is set correctly
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Roboto')
        
        run.font.size = Pt(12)  # Adjust the size as needed
        run.font.color.rgb = RGBColor(0, 51, 153)  # Dark blue color
        # set invoice No
        row = table.rows[3].cells
        cell = row[4]
        # Set text properties
        paragraph = cell.paragraphs[0]
        paragraph.clear()  # Clear existing text in the paragraph
        run = paragraph.add_run('#' + str(invoice.id))
        run.bold = True
        run.font.name = 'Roboto'

        # This part ensures the font name is set correctly
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Roboto')
        
        run.font.size = Pt(12)  # Adjust the size as needed
        run.font.color.rgb = RGBColor(0, 51, 153)  # Dark blue color
    
    def handle_table1(table):
        row = table.rows[2].cells
        row[0].text = ""+ str(invoice.name.name)
        
        # Company Name
        row = table.rows[3].cells
        row[0].text = ""+ str(invoice.name.company)
        # Address
        row = table.rows[4].cells
        row[0].text = ""+ str(invoice.name.city)
        # Phone
        row = table.rows[5].cells
        row[0].text = ""+ str(invoice.name.contact_number)
    
    def handle_table2(table):
        total = 0
        row_index = 1  # Assuming the first row is the header row

        for item in client_data:
            row = table.add_row()
            # If there are not enough rows in the table, you can decide to add new rows if necessary
            row = table.rows[row_index].cells
            if float(item.get('quantity')) > 0.0:
                row[0].text = str(item.get('name'))
                
                row[2].text = str(item.get('quantity'))
                row[3].text = str(item.get('price'))
                
                price = int(item.get('quantity', 0)) * float(item.get('price', 0))
                row[4].text = str(price)
                total += price
                row_index += 1

    def handle_table3(table):
        total = 0
        row_index = 1  # Assuming the first row is the header row
        total = 1000  # Dummy total value, replace with actual calculation if available

        # Starting row index (assuming the first row is the header)
        start_row_index = 6  # Adjust this based on where you want to start

        max_length = max(len(installments_list), len(final_block_data))

        for index in range(max_length):
            row = table.add_row().cells
            installment = list_get(installments_list, index, default={})
            label, value = list_get(final_block_data, index, default=("", ""))

            # Add installment data if available
            
            row[0].text = str(installment.get('amount', ''))
            
            formatted_date = format_date(installment.get('time', ''))
            row[1].text = formatted_date
            for paragraph in row[1].paragraphs:
                for run in paragraph.runs:
                    set_font_style(run, 'Roboto', Pt(10))
            # Add final block data if available and not over index limit
            if index < len(final_block_data):
                row[3].text = label
                row[4].text = str(value)
                if label == "Total":
                    set_cell_background(row[4], "00ff00")
                else:
                    set_cell_background(row[4], "eeeeee")

            row_index += 1

    doc = Document(file_path)
    
    table = doc.tables[0]
    handle_table0(table)
    
    # Adding Client data
    table = doc.tables[1]
    handle_table1(table)
    # Set the default font and size for the entire document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Roboto'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

    # Additional settings for font, especially for specific scripts
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Roboto')
    table = doc.tables[2]
    handle_table2(table)
    
    table = doc.tables[3]
    handle_table3(table)
    tmp = tempfile.NamedTemporaryFile(delete=False)
    
    # doc.save(tmp.name)
    tmp.close()  # Manually close the file to ensure all data is written
    from docxcompose.composer import Composer
    
    header = Document(header_path)
    def replace_paragraph_text(doc, index, new_text):
        if index < len(doc.paragraphs):
            paragraph = doc.paragraphs[index]
            paragraph.clear()  # Clear existing text
            run = paragraph.add_run(new_text)  # Add new text
            add_styling(run)
        else:
            print(f"No paragraph found at index {index}")

    # Replace text in paragraphs 2 and 3
    replace_paragraph_text(header, 2, str(invoice.system_capacity) + ' KW')
    replace_paragraph_text(header, 3, str(invoice.name.name))
    composer = Composer(header)
    footer = Document(footer_path)
    composer.append(doc)
    composer.append(footer)
    composer.save(tmp.name)
    tmp = open(tmp.name, 'rb')
    response = FileResponse(tmp, as_attachment=True, filename=f'modified_invoice_{invoice.id}.docx')
    response.headers['Content-Disposition'] = f'attachment; filename="modified_invoice_{invoice.id}.docx"'
    return response
class GenericModelListView(ListView):
    template_name = 'generic_list.html'
    action = 'list'
    paginate_by = 5
    models = ['Client', 'SolarPanel', 'Inverter', 'Structure', 'Cabling', 'NetMetering', 'Batteries','LightningArrestor','Installation','Invoice']
    
    def get_queryset(self,request=None):
        self.model = apps.get_model('crud', self.kwargs['model_name'])
        queryset = self.model.objects.all()
        queryset = self.apply_filters(queryset)
        queryset = self.apply_sort(queryset)
        return queryset
    
    def apply_filters(self, queryset):
        search_query = self.request.GET.get('q')
        date_range = self.request.GET.get('date_range')
        if search_query:
            query = Q()
            for field in self.get_searchable_fields():
                query |= Q(**{f"{field}__icontains": search_query})
            queryset = queryset.filter(query)
        if date_range:
            try:
                start_date, end_date = date_range.split(' to ')
                start_date = datetime.strptime(start_date, '%Y-%m-%d')
                end_date = datetime.strptime(end_date, '%Y-%m-%d')
                queryset = queryset.filter(created_at__range=(start_date, end_date))
            except ValueError:
                pass
        return queryset

    def get_searchable_fields(self):
        fields = []
        for field in self.model._meta.fields:
            if isinstance(field, (models.CharField, models.TextField)):
                fields.append(field.name)
            elif isinstance(field, models.ForeignKey):
                # Adjust the search to use 'name' or 'brand' for ForeignKey fields
                if hasattr(field.related_model, 'name'):
                    fields.append(f"{field.name}__name")
                elif hasattr(field.related_model, 'brand'):
                    fields.append(f"{field.name}__brand")
        return fields

    def apply_sort(self, queryset):
        sort_by = self.request.GET.get('sort_by', None)
        if sort_by:
            if sort_by.startswith('-'):
                sort_by = sort_by[1:]
            else:
                sort_by = f'-{sort_by}'
            queryset = queryset.order_by(sort_by)
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['model_name'] = self.kwargs['model_name']
        context['field_labels'] = self.model.get_field_labels()
        return context

class GenericModelCreateView(CreateView):

    action = 'ADD'
    def get_template_names(self):
        model_name = self.kwargs.get('model_name')
        if model_name == 'Invoice':
            return ['invoice_form.html']  # Template name for Invoice model
        else:
            return ['generic_form.html']  # Template name for other models
    

    def get_form_class(self):
        model_name = self.kwargs['model_name']
        form_class_name = self.kwargs['model_name'] + 'Form'
        print('form class name is --> ',form_class_name)
        return getattr(forms, form_class_name)
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method in ('POST', 'PUT'):
            kwargs.update({'files': self.request.FILES})
        return kwargs
    def get_context_data(self, **kwargs):
        context = super(GenericModelCreateView, self).get_context_data(**kwargs)
        context['model_name'] = self.kwargs.get('model_name', 'InvoiceModel')
        context['action'] = self.action
        return context
    def get_success_url(self):
        return reverse_lazy('generic_list', kwargs={'model_name': self.kwargs['model_name']})

class GenericModelUpdateView(UpdateView):
    
    action = 'update'
    def get_template_names(self):
        model_name = self.kwargs.get('model_name')
        if model_name == 'Invoice':
            return ['invoice_form.html']  # Template name for Invoice model
        else:
            return ['generic_form.html']  # Template name for other models

    def get_object(self, queryset=None):
        model = apps.get_model('crud', self.kwargs['model_name'])
        return model.objects.get(pk=self.kwargs['pk'])

    def get_form_class(self):
        form_class_name = self.kwargs['model_name'] + 'Form'
        return getattr(forms, form_class_name)
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method in ('POST', 'PUT'):
            kwargs.update({'files': self.request.FILES})
        return kwargs
    def get_context_data(self, **kwargs):
        context = super(GenericModelUpdateView, self).get_context_data(**kwargs)
        context['model_name'] = self.kwargs.get('model_name', 'InvoiceModel')
        context['action'] = self.action
        return context
    def get_success_url(self):  
        return reverse_lazy('generic_list', kwargs={'model_name': self.kwargs['model_name']})

class GenericModelDeleteView(DeleteView):
    template_name = 'generic_confirm_delete.html'
    action = 'delete'
    def get_context_data(self, **kwargs):
        context = super(GenericModelDeleteView, self).get_context_data(**kwargs)
        context['model_name'] = self.kwargs.get('model_name', 'InvoiceModel')
        context['action'] = self.action
        return context
    def get_object(self, queryset=None):
        model = apps.get_model('crud', self.kwargs['model_name'])
        return model.objects.get(pk=self.kwargs['pk'])

    def get_success_url(self):
        return reverse_lazy('generic_list', kwargs={'model_name': self.kwargs['model_name']})

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['model_name'] = self.kwargs['model_name']
        return context

    """
    A ViewSet that determines the queryset and serializer class dynamically
    based on the 'model_name' URL keyword argument.
    """

    def get_queryset(self):
        model_name = self.kwargs.get('model_name')
        if model_name:
            model = apps.get_model('crud', model_name)
            return model.objects.all()
        else:
            raise ImproperlyConfigured("No model specified!")

    def get_serializer_class(self):
        model_name = self.kwargs.get('model_name')
        if model_name:
            class Meta:
                model = apps.get_model('crud', model_name)
                model = model
                fields = '__all__'

            serializer_class = type(
                f"{model_name}Serializer",
                (serializers.ModelSerializer,),
                {'Meta': Meta}
            )
            return serializer_class
        else:
            raise ImproperlyConfigured("No model specified!")


class InvoiceModelListView(APIView):
    def get_queryset(self, model_name, invoice_id=None):
        self.model = apps.get_model('crud', model_name)
        if invoice_id:
            queryset = self.model.objects.filter(id=invoice_id)
        else:
            queryset = self.model.objects.all()
            queryset = self.apply_filters(queryset)
            queryset = self.apply_sort(queryset)
        return queryset
    
    def apply_filters(self, queryset):
        search_query = self.request.GET.get('q')
        date_range = self.request.GET.get('date_range')
        if search_query:
            query = Q()
            for field in self.get_searchable_fields():
                query |= Q(**{f"{field}__icontains": search_query})
            queryset = queryset.filter(query)
        if date_range:
            try:
                start_date, end_date = date_range.split(' to ')
                start_date = datetime.strptime(start_date, '%Y-%m-%d')
                end_date = datetime.strptime(end_date, '%Y-%m-%d')
                queryset = queryset.filter(created_at__range=(start_date, end_date))
            except ValueError:
                pass
        return queryset

    def get_searchable_fields(self):
        fields = []
        for field in self.model._meta.fields:
            if isinstance(field, (models.CharField, models.TextField)):
                fields.append(field.name)
            elif isinstance(field, models.ForeignKey):
                if hasattr(field.related_model, 'name'):
                    fields.append(f"{field.name}__name")
                elif hasattr(field.related_model, 'brand'):
                    fields.append(f"{field.name}__brand")
        return fields

    def apply_sort(self, queryset):
        sort_by = self.request.GET.get('sort_by', None)
        if sort_by:
            queryset = queryset.order_by(sort_by)
        return queryset
    
    def get_serializer_class(self, model_name):
        module_path = 'crud.serializers'
        serializer_class_name = model_name + 'Serializer'

        try:
            serializers_module = importlib.import_module(module_path)
            serializer_class = getattr(serializers_module, serializer_class_name)
        except (ImportError, AttributeError):
            raise NotFound(f"Serializer for model '{model_name}' not found.")
        
        return serializer_class

    def get_related_data(self, instance):
        related_data = {}
        related_fields = [
            'name', 'solar_panel', 'inverter', 'structure', 
            'cabling', 'net_metering', 'battery', 
            'lightning_arrestor', 'installation'
        ]
        
        for field in related_fields:
            related_instance = getattr(instance, field)
            if related_instance:
                serializer_class = self.get_serializer_class(related_instance.__class__.__name__)
                related_data[field] = serializer_class(related_instance).data
            else:
                related_data[field] = None
        
        return related_data

    def get(self, request, *args, **kwargs):
        model_name = 'Invoice'
        invoice_id = kwargs.get('invoice_id')  # Extract invoice_id from the URL kwargs
        queryset = self.get_queryset(model_name, invoice_id=invoice_id)
        serializer_class = self.get_serializer_class(model_name)
        serializer = serializer_class(queryset, many=True)
        
        detailed_data = []
        for instance in queryset:
            instance_data = serializer_class(instance).data
            instance_data.update(self.get_related_data(instance))
            # Fetch related data from PartiallyPaid table based on invoice_id
            partial_payments = PartiallyPaid.objects.filter(invoice_id=instance.id)
            partial_payment_data = []
            for payment in partial_payments:
                partial_payment_data.append({
                    'client_id': payment.client_id,
                    'amount': payment.amount,
                    'time': payment.time
                })
            
            # Include the partial payment data in the response
            instance_data['partial_payments'] = partial_payment_data
            detailed_data.append(instance_data)
        
        return Response(detailed_data, status=status.HTTP_200_OK)

    def post(self, request, *args, **kwargs):
        model_name = 'Invoice'
        serializer_class = self.get_serializer_class(model_name)
        serializer = serializer_class(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
